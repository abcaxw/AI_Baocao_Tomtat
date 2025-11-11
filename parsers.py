"""
Document Parsers with OCR Support - Đọc PDF, DOCX, XLSX
Supports both text-based and scanned (image-based) PDFs
"""

import PyPDF2
from docx import Document
import pandas as pd
from pathlib import Path
from typing import Optional
import os


class DocumentParser:
    """Parse các loại file khác nhau, bao gồm OCR cho PDF scan"""

    def __init__(self):
        self.ocr_available = self._check_ocr_availability()

    def _check_ocr_availability(self) -> bool:
        """Check if OCR libraries are available"""
        try:
            import pytesseract
            from pdf2image import convert_from_path
            return True
        except ImportError:
            return False

    def parse(self, file_path: str) -> str:
        """Parse file và trả về text content"""
        file_path = Path(file_path)
        ext = file_path.suffix.lower()

        parsers = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.xlsx': self._parse_xlsx,
            '.xls': self._parse_xlsx,
            '.txt': self._parse_txt
        }

        parser_func = parsers.get(ext)
        if not parser_func:
            raise ValueError(f"Unsupported file type: {ext}")

        content = parser_func(file_path)

        if not content or len(content.strip()) == 0:
            raise ValueError(
                f"No text content could be extracted from {file_path.name}. "
                "If this is a scanned PDF, make sure OCR libraries are installed."
            )

        return content

    def _parse_pdf(self, file_path: Path) -> str:
        """Parse PDF file with OCR fallback"""
        print(f"📄 Attempting to parse PDF: {file_path.name}")

        # Try standard text extraction first
        text = self._parse_pdf_text(file_path)

        if text and len(text.strip()) > 100:  # If we got meaningful text
            print(f"✅ Extracted {len(text)} characters using standard method")
            return text.strip()

        # If standard extraction failed, try OCR
        print("⚠️  Standard extraction failed or insufficient text")

        if self.ocr_available:
            print("🔍 Attempting OCR (this may take a while for large PDFs)...")
            text = self._parse_pdf_with_ocr(file_path)
            if text:
                print(f"✅ OCR extracted {len(text)} characters")
                return text.strip()
        else:
            print("❌ OCR not available. Install with:")
            print("   pip install pytesseract pdf2image")
            print("   Also install Tesseract: https://github.com/tesseract-ocr/tesseract")

        raise Exception("Could not extract text from PDF using any method")

    def _parse_pdf_text(self, file_path: Path) -> str:
        """Try to extract text from PDF (non-OCR)"""
        text = ""

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                print(f"   📄 PDF has {len(pdf_reader.pages)} pages")

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and len(page_text.strip()) > 0:
                            text += page_text + "\n"
                    except:
                        continue

            return text

        except Exception as e:
            print(f"   Error in standard extraction: {str(e)}")
            return ""

    def _parse_pdf_with_ocr(self, file_path: Path) -> str:
        """Parse PDF using OCR for scanned documents"""
        try:
            import pytesseract
            from pdf2image import convert_from_path
            from PIL import Image

            # Configure tesseract path (Windows - adjust if needed)
            if os.name == 'nt':  # Windows
                tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
                if os.path.exists(tesseract_path):
                    pytesseract.pytesseract.tesseract_cmd = tesseract_path

            # Convert PDF to images
            print("   Converting PDF pages to images...")

            # For large PDFs, process in batches
            images = convert_from_path(
                file_path,
                dpi=200,  # Lower DPI for faster processing (300 is better quality)
                thread_count=4
            )

            print(f"   Processing {len(images)} pages with OCR...")

            text = ""
            for i, image in enumerate(images):
                print(f"   OCR Page {i+1}/{len(images)}...", end='\r')

                # Run OCR on image
                page_text = pytesseract.image_to_string(image, lang='eng')
                text += page_text + "\n\n"

            print()  # New line after progress
            return text

        except ImportError as e:
            raise Exception(
                f"OCR libraries not installed: {str(e)}\n"
                "Install with: pip install pytesseract pdf2image pillow\n"
                "Also install Tesseract-OCR: https://github.com/tesseract-ocr/tesseract"
            )
        except Exception as e:
            raise Exception(f"OCR failed: {str(e)}")

    def _parse_docx(self, file_path: Path) -> str:
        """Parse DOCX file"""
        try:
            doc = Document(file_path)
            text = ""

            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"

            if not text.strip():
                raise ValueError("No text content found in DOCX")

            return text.strip()
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")

    def _parse_xlsx(self, file_path: Path) -> str:
        """Parse XLSX/XLS file"""
        try:
            excel_file = pd.ExcelFile(file_path)
            text = ""

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                if df.empty:
                    continue

                text += f"\n=== Sheet: {sheet_name} ===\n"
                text += df.to_string(index=False)
                text += "\n"

            if not text.strip():
                raise ValueError("No data found in Excel file")

            return text.strip()
        except Exception as e:
            raise Exception(f"Error parsing Excel: {str(e)}")

    def _parse_txt(self, file_path: Path) -> str:
        """Parse TXT file"""
        for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read().strip()
                    if content:
                        return content
            except:
                continue

        raise Exception("Could not decode text file")

    def get_metadata(self, file_path: str) -> dict:
        """Lấy metadata của file"""
        file_path = Path(file_path)
        return {
            "filename": file_path.name,
            "extension": file_path.suffix,
            "size_bytes": file_path.stat().st_size,
            "size_mb": round(file_path.stat().st_size / (1024 * 1024), 2),
            "ocr_available": self.ocr_available
        }


def test_parser(file_path: str):
    """Test parser with a file"""
    parser = DocumentParser()

    print("=" * 60)
    print(f"Testing parser with: {file_path}")
    print("=" * 60)

    try:
        metadata = parser.get_metadata(file_path)
        print(f"\n📊 File Info:")
        print(f"   Name: {metadata['filename']}")
        print(f"   Type: {metadata['extension']}")
        print(f"   Size: {metadata['size_mb']} MB")
        print(f"   OCR Available: {'✅ Yes' if metadata['ocr_available'] else '❌ No'}")
        print()

        content = parser.parse(file_path)

        print(f"\n✅ Successfully extracted {len(content)} characters")
        print(f"\n📄 Preview (first 800 chars):")
        print("-" * 60)
        print(content[:800])
        print("-" * 60)

        # Save to file for inspection
        output_file = f"extracted_{Path(file_path).stem}.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"\n💾 Full content saved to: {output_file}")

        return content

    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        test_parser(sys.argv[1])
    else:
        print("Usage: python parsers.py <file_path>")
        print("Example: python parsers.py document.pdf")